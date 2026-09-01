from pathlib import Path
import csv
import io
import re
from typing import Any

from pypdf import PdfReader
from docx import Document as DocxDocument
from openpyxl import load_workbook


class DocumentProcessor:
    """
    Converts knowledge-base documents into meaningful chunks.

    Chunking strategy:
        Markdown:
            Split by headings and keep sections together.

        Python:
            Split primarily around classes/functions/methods.

        PDF / Word / plain text:
            Split by paragraphs (PDF pages become paragraph
            boundaries).

        CSV / Excel:
            Each row becomes a small readable text block.

    Large sections are split only when necessary.
    """

    SUPPORTED_EXTENSIONS = {
        ".txt",
        ".md",
        ".py",
        ".pdf",
        ".docx",
        ".csv",
        ".xlsx",
    }

    STRUCTURED_EXTENSIONS = {
        ".pdf",
        ".docx",
        ".csv",
        ".xlsx",
    }

    def __init__(
        self,
        max_chunk_characters: int = 6000,
        overlap_characters: int = 500,
    ):
        self.max_chunk_characters = max(
            1000,
            max_chunk_characters,
        )

        self.overlap_characters = max(
            0,
            min(
                overlap_characters,
                self.max_chunk_characters // 2,
            ),
        )

    # ============================================================
    # PUBLIC API
    # ============================================================

    def process(
        self,
        file_path: str,
    ) -> list[dict[str, Any]]:

        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(
                f"Document not found: {file_path}"
            )

        extension = path.suffix.lower()

        if extension not in self.SUPPORTED_EXTENSIONS:
            return []

        if extension in self.STRUCTURED_EXTENSIONS:
            text = self._extract_structured_text(
                path, extension
            )
        else:
            text = path.read_text(
                encoding="utf-8",
                errors="ignore",
            ).strip()

        if not text:
            return []

        if extension == ".md":
            sections = self._split_markdown(text)

        elif extension == ".py":
            sections = self._split_python(text)

        else:
            # .txt, .pdf, .docx, .csv, .xlsx all end up as
            # paragraph-separated plain text by this point.
            sections = self._split_text(text)

        chunks = []

        for section in sections:

            section = section.strip()

            if not section:
                continue

            if len(section) <= self.max_chunk_characters:
                chunks.append(section)
            else:
                chunks.extend(
                    self._split_large_section(section)
                )

        return chunks

    # ============================================================
    # BINARY FORMAT TEXT EXTRACTION
    # ============================================================

    def _extract_structured_text(
        self,
        path: Path,
        extension: str,
    ) -> str:

        if extension == ".pdf":
            return self._extract_pdf_text(path)

        if extension == ".docx":
            return self._extract_docx_text(path)

        if extension == ".xlsx":
            return self._extract_xlsx_text(path)

        if extension == ".csv":
            return self._extract_csv_text(path)

        return ""

    def _extract_pdf_text(self, path: Path) -> str:

        reader = PdfReader(str(path))

        pages = []

        for page in reader.pages:

            page_text = (page.extract_text() or "").strip()

            if page_text:
                pages.append(page_text)

        # Blank line between pages so paragraph-splitting
        # treats each page as its own section boundary.
        return "\n\n".join(pages).strip()

    def _extract_docx_text(self, path: Path) -> str:

        document = DocxDocument(str(path))

        paragraphs = [
            paragraph.text.strip()
            for paragraph in document.paragraphs
            if paragraph.text.strip()
        ]

        # Tables aren't in .paragraphs — pull their text too.
        for table in document.tables:
            for row in table.rows:

                cells = [
                    cell.text.strip()
                    for cell in row.cells
                    if cell.text.strip()
                ]

                if cells:
                    paragraphs.append(" | ".join(cells))

        return "\n\n".join(paragraphs).strip()

    def _extract_xlsx_text(self, path: Path) -> str:

        workbook = load_workbook(
            str(path),
            read_only=True,
            data_only=True,
        )

        blocks = []

        for sheet in workbook.worksheets:

            header = None

            for row in sheet.iter_rows(values_only=True):

                values = [
                    str(cell) for cell in row if cell is not None
                ]

                if not values:
                    continue

                if header is None:
                    header = values
                    continue

                pairs = [
                    f"{key}: {value}"
                    for key, value in zip(header, values)
                ]

                if pairs:
                    blocks.append(
                        f"[{sheet.title}] "
                        + ", ".join(pairs)
                    )

        return "\n\n".join(blocks).strip()

    def _extract_csv_text(self, path: Path) -> str:

        with open(
            path, "r", encoding="utf-8", errors="ignore"
        ) as csv_file:

            reader = csv.reader(csv_file)
            rows = list(reader)

        if not rows:
            return ""

        header = rows[0]

        blocks = []

        for row in rows[1:]:

            pairs = [
                f"{key}: {value}"
                for key, value in zip(header, row)
                if value
            ]

            if pairs:
                blocks.append(", ".join(pairs))

        return "\n\n".join(blocks).strip()

    # ============================================================
    # MARKDOWN CHUNKING
    # ============================================================

    def _split_markdown(
        self,
        text: str,
    ) -> list[str]:

        lines = text.splitlines()

        sections = []
        current = []

        for line in lines:

            # Markdown heading
            if re.match(
                r"^\s{0,3}#{1,6}\s+",
                line,
            ):

                if current:
                    sections.append(
                        "\n".join(current).strip()
                    )
                    current = []

            current.append(line)

        if current:
            sections.append(
                "\n".join(current).strip()
            )

        return [
            section
            for section in sections
            if section.strip()
        ]

    # ============================================================
    # PYTHON CHUNKING
    # ============================================================

    def _split_python(
        self,
        text: str,
    ) -> list[str]:

        lines = text.splitlines()

        sections = []
        current = []

        for index, line in enumerate(lines):

            stripped = line.lstrip()

            is_top_level_definition = (
                (
                    stripped.startswith("class ")
                    or stripped.startswith("def ")
                    or stripped.startswith("async def ")
                )
                and len(line) - len(stripped) == 0
            )

            if is_top_level_definition and current:

                sections.append(
                    "\n".join(current).strip()
                )

                current = []

            current.append(line)

        if current:
            sections.append(
                "\n".join(current).strip()
            )

        return [
            section
            for section in sections
            if section.strip()
        ]

    # ============================================================
    # PLAIN TEXT CHUNKING
    # ============================================================

    def _split_text(
        self,
        text: str,
    ) -> list[str]:

        paragraphs = re.split(
            r"\n\s*\n",
            text,
        )

        return [
            paragraph.strip()
            for paragraph in paragraphs
            if paragraph.strip()
        ]

    # ============================================================
    # LARGE SECTION SPLITTING
    # ============================================================

    def _split_large_section(
        self,
        text: str,
    ) -> list[str]:

        chunks = []

        # First try paragraph boundaries.
        paragraphs = re.split(
            r"\n\s*\n",
            text,
        )

        current = ""

        for paragraph in paragraphs:

            paragraph = paragraph.strip()

            if not paragraph:
                continue

            candidate = (
                f"{current}\n\n{paragraph}"
                if current
                else paragraph
            )

            if (
                len(candidate)
                <= self.max_chunk_characters
            ):
                current = candidate
                continue

            if current:
                chunks.append(current)

            # A single paragraph is still too large.
            if len(paragraph) > self.max_chunk_characters:

                sub_chunks = self._split_by_lines(
                    paragraph
                )

                chunks.extend(sub_chunks)

                current = ""

            else:
                current = paragraph

        if current:
            chunks.append(current)

        return chunks

    # ============================================================
    # LINE-BASED FALLBACK
    # ============================================================

    def _split_by_lines(
        self,
        text: str,
    ) -> list[str]:

        lines = text.splitlines()

        chunks = []
        current_lines = []
        current_length = 0

        for line in lines:

            line_length = len(line) + 1

            if (
                current_lines
                and current_length + line_length
                > self.max_chunk_characters
            ):

                chunk = "\n".join(
                    current_lines
                ).strip()

                if chunk:
                    chunks.append(chunk)

                # Small overlap from previous lines.
                overlap_lines = []

                overlap_length = 0

                for previous_line in reversed(
                    current_lines
                ):

                    if (
                        overlap_length
                        + len(previous_line)
                        + 1
                        > self.overlap_characters
                    ):
                        break

                    overlap_lines.insert(
                        0,
                        previous_line,
                    )

                    overlap_length += (
                        len(previous_line) + 1
                    )

                current_lines = overlap_lines
                current_length = overlap_length

            current_lines.append(line)
            current_length += line_length

        if current_lines:

            chunk = "\n".join(
                current_lines
            ).strip()

            if chunk:
                chunks.append(chunk)

        return chunks


# ============================================================
# SHARED INSTANCE
# ============================================================

document_processor = DocumentProcessor(
    max_chunk_characters=6000,
    overlap_characters=500,
)